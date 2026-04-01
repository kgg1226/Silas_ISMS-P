"""
Word 문서(.docx) 파서 — python-docx 사용
Heading 스타일 기반 섹션 분리 + 본문 텍스트 추출.
"""

from __future__ import annotations

from pathlib import Path

from parsers.section_detector import RawSection


def parse_docx(file_path: str | Path) -> tuple[list[RawSection], int]:
    """
    Word 문서에서 텍스트 추출.
    Heading 단락은 별도 RawSection으로 분리.

    Returns:
        (raw_sections, estimated_pages)
        Word는 정확한 페이지 수를 알 수 없으므로 추정값 반환.
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError(
            "python-docx가 설치되지 않았습니다. pip install python-docx"
        )

    doc = Document(str(file_path))
    raw_sections: list[RawSection] = []
    current_text_lines: list[str] = []
    page_estimate = 1
    line_count = 0

    def _flush():
        nonlocal page_estimate
        if current_text_lines:
            text = "\n".join(current_text_lines)
            raw_sections.append(RawSection(
                text=text,
                page_number=page_estimate,
            ))
            current_text_lines.clear()

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        line_count += 1
        # 대략 40줄당 1페이지 추정
        if line_count % 40 == 0:
            page_estimate += 1

        # Heading 스타일 감지
        style_name = (para.style.name or "").lower()
        if "heading" in style_name or "제목" in style_name:
            _flush()
            # Heading은 별도 RawSection으로
            raw_sections.append(RawSection(
                text=text,
                page_number=page_estimate,
            ))
        else:
            current_text_lines.append(text)

    _flush()

    return raw_sections, page_estimate


def parse_docx_simple(file_path: str | Path) -> tuple[list[RawSection], int]:
    """
    간단 모드: 전체 텍스트를 하나의 RawSection으로.
    Heading 구분 없이 전체 텍스트 추출.
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx가 설치되지 않았습니다.")

    doc = Document(str(file_path))
    full_text = "\n".join(
        para.text.strip() for para in doc.paragraphs if para.text.strip()
    )

    if not full_text:
        return [], 0

    # 추정 페이지 수
    lines = full_text.count("\n") + 1
    pages = max(1, lines // 40)

    return [RawSection(text=full_text, page_number=1)], pages
